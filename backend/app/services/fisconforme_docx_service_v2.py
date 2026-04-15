from __future__ import annotations

import base64
import io
from datetime import datetime
from pathlib import Path

from backend.app.config import settings
from backend.app.services.fisconforme_dsf_read_service_v2 import read_dsf_pdf_base64_v2
from pipeline.fisconforme.query_service_v2 import read_fisconforme_cache_v2


def _generated_root() -> Path:
    path = settings.app_state_root / "fisconforme_v2" / "generated"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _safe_output_dir(output_dir: str) -> Path | None:
    if not output_dir.strip() or ".." in output_dir:
        return None
    target = (settings.workspace_root / output_dir.lstrip("/\\")).resolve()
    if not target.is_relative_to(settings.workspace_root.resolve()):
        return None
    target.mkdir(parents=True, exist_ok=True)
    return target


def _resolve_pdf_base64(payload: dict) -> str:
    explicit = str(payload.get("pdf_base64", "") or "")
    if explicit:
        return explicit
    dsf_id = str(payload.get("dsf_id", "") or "")
    if not dsf_id:
        return ""
    return read_dsf_pdf_base64_v2(dsf_id)


def _pdf_base64_to_png_streams(pdf_base64: str) -> list[io.BytesIO]:
    if not pdf_base64:
        return []
    try:
        import fitz
    except ImportError:
        return []
    try:
        pdf_bytes = base64.b64decode(pdf_base64)
        doc = fitz.open(stream=io.BytesIO(pdf_bytes), filetype="pdf")
        streams: list[io.BytesIO] = []
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            stream = io.BytesIO(pix.tobytes("png"))
            stream.seek(0)
            streams.append(stream)
        doc.close()
        return streams
    except Exception:
        return []


def build_notification_docx_v2(cnpj: str, payload: dict) -> dict:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("python-docx não instalado no ambiente atual") from exc

    overview = read_fisconforme_cache_v2(cnpj)
    cadastral = list(overview.get("dados_cadastrais", []) or [])
    malhas = list(overview.get("malhas", []) or [])
    first = cadastral[0] if cadastral else {}
    razao_social = str(first.get("razao_social", "") or first.get("nome", "") or "")
    ie = str(first.get("ie", "") or "")

    document = Document()
    document.add_heading("NOTIFICAÇÃO FISCONFORME NÃO ATENDIDO", level=1)
    document.add_paragraph(f"Razão Social: {razao_social}")
    document.add_paragraph(f"CNPJ: {cnpj}")
    document.add_paragraph(f"IE: {ie}")
    document.add_paragraph(f"DSF: {str(payload.get('dsf', '') or '')}")
    document.add_paragraph(f"Auditor: {str(payload.get('auditor', '') or '')}")
    document.add_paragraph(f"Cargo/Título: {str(payload.get('cargo_titulo', '') or '')}")
    document.add_paragraph(f"Matrícula: {str(payload.get('matricula', '') or '')}")
    document.add_paragraph(f"Contato: {str(payload.get('contato', '') or '')}")
    document.add_paragraph(f"Órgão de origem: {str(payload.get('orgao_origem', '') or '')}")

    document.add_heading("Pendências", level=2)
    if malhas:
        cols = ["id_pendencia", "id_notificacao", "titulo_malha", "periodo", "status_pendencia"]
        labels = ["ID Pend.", "ID Notif.", "Título", "Período", "Status"]
        table = document.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for idx, label in enumerate(labels):
            hdr[idx].text = label
        for row in malhas:
            cells = table.add_row().cells
            for idx, key in enumerate(cols):
                cells[idx].text = str(row.get(key, "") or "")
    else:
        document.add_paragraph("(Sem pendências registradas)")

    image_streams = _pdf_base64_to_png_streams(_resolve_pdf_base64(payload))
    if image_streams:
        document.add_heading("Imagens da DSF", level=2)
        for stream in image_streams:
            document.add_picture(stream, width=Inches(5.8))

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"notificacao_det_{cnpj}_{timestamp}.docx"
    base_dir = _generated_root()
    path = base_dir / file_name
    document.save(str(path))

    extra_saved = ""
    output_dir = str(payload.get("output_dir", "") or "")
    target_dir = _safe_output_dir(output_dir) if output_dir else None
    if target_dir is not None:
        final_path = target_dir / file_name
        final_path.write_bytes(path.read_bytes())
        extra_saved = str(final_path)

    return {
        "cnpj": cnpj,
        "nome_arquivo": file_name,
        "docx_path": str(path),
        "salvo_em": extra_saved,
        "malhas_count": len(malhas),
        "tem_imagens_dsf": bool(image_streams),
    }
