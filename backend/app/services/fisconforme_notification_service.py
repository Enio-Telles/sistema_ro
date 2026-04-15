from __future__ import annotations

import base64
import json
import re
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from backend.app.config import settings
from pipeline.fisconforme.cache import load_cache

try:  # pragma: no cover - optional dependency
    import fitz  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    fitz = None

try:  # pragma: no cover - optional dependency
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    Document = None
    Inches = None


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _template_path() -> Path:
    return _repo_root() / "modelo" / "modelo_notificacao_fisconforme_n_atendido.txt"


def _state_root() -> Path:
    path = settings.app_state_root / "fisconforme"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _auditor_config_path() -> Path:
    return _state_root() / "auditor_config.json"


def sanitize_cnpj(cnpj: str) -> str:
    return re.sub(r"\D", "", cnpj or "")


def validate_cnpj(cnpj: str) -> bool:
    normalized = sanitize_cnpj(cnpj)
    if len(normalized) != 14 or len(set(normalized)) == 1:
        return False

    pesos1 = [5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    soma1 = sum(int(normalized[i]) * pesos1[i] for i in range(12))
    resto1 = soma1 % 11
    digito1 = 0 if resto1 < 2 else 11 - resto1

    pesos2 = [6, 5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    soma2 = sum(int(normalized[i]) * pesos2[i] for i in range(13))
    resto2 = soma2 % 11
    digito2 = 0 if resto2 < 2 else 11 - resto2

    return digito1 == int(normalized[12]) and digito2 == int(normalized[13])


def load_auditor_config() -> dict[str, str]:
    path = _auditor_config_path()
    if not path.exists():
        return {
            "auditor": "",
            "cargo_titulo": "",
            "matricula": "",
            "contato": "",
            "orgao_origem": "",
        }

    payload = json.loads(path.read_text(encoding="utf-8"))
    return {
        "auditor": str(payload.get("auditor", "") or ""),
        "cargo_titulo": str(payload.get("cargo_titulo", "") or ""),
        "matricula": str(payload.get("matricula", "") or ""),
        "contato": str(payload.get("contato", "") or ""),
        "orgao_origem": str(payload.get("orgao_origem", "") or ""),
    }


def save_auditor_config(
    *,
    auditor: str,
    cargo_titulo: str = "",
    matricula: str = "",
    contato: str = "",
    orgao_origem: str = "",
) -> Path:
    payload = {
        "auditor": auditor.strip(),
        "cargo_titulo": cargo_titulo.strip(),
        "matricula": matricula.strip(),
        "contato": contato.strip(),
        "orgao_origem": orgao_origem.strip(),
        "saved_at": datetime.now().isoformat(),
    }
    path = _auditor_config_path()
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _load_cadastral_record(cnpj: str) -> dict[str, Any]:
    df = load_cache(cnpj, "fisconforme_cadastral")
    if df is None or df.is_empty():
        return {}
    return df.to_dicts()[0]


def _load_malhas_records(cnpj: str) -> list[dict[str, Any]]:
    df = load_cache(cnpj, "fisconforme_malhas")
    if df is None or df.is_empty():
        return []
    return df.to_dicts()


def _pick_value(record: dict[str, Any], *keys: str) -> str:
    for key in keys:
        for candidate in (key, key.upper(), key.lower()):
            value = record.get(candidate)
            if value not in (None, ""):
                return str(value)
    return ""


def render_malhas_html_table(malhas: list[dict[str, Any]]) -> str:
    if not malhas:
        return "<p><em>(Sem pendências registradas)</em></p>"

    columns = [
        ("ID Pend.", "id_pendencia"),
        ("ID Notif.", "id_notificacao"),
        ("Malha ID", "malhas_id"),
        ("Título", "titulo_malha"),
        ("Período", "periodo"),
        ("Status Pend.", "status_pendencia"),
        ("Status Notif.", "status_notificacao"),
        ("Ciência", "data_ciencia_consolidada"),
    ]
    parts = [
        "<table style='border-collapse:collapse;width:100%;'>",
        "<thead><tr>",
    ]
    for label, _ in columns:
        parts.append(
            "<th style='border:1px solid #ccc;padding:4px 8px;background:#f0f0f0;"
            "font-family:Arial,sans-serif;font-size:10px;'>"
            f"{label}</th>"
        )
    parts.append("</tr></thead><tbody>")

    for malha in malhas:
        parts.append("<tr>")
        for _, key in columns:
            value = _pick_value(malha, key)
            parts.append(
                "<td style='border:1px solid #ccc;padding:4px 8px;"
                "font-family:Arial,sans-serif;font-size:10px;'>"
                f"{value}</td>"
            )
        parts.append("</tr>")

    parts.append("</tbody></table>")
    return "".join(parts)


def pdf_base64_to_png_pages(pdf_base64: str | None) -> list[bytes]:
    if not pdf_base64 or fitz is None:
        return []

    pdf_bytes = base64.b64decode(pdf_base64)
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages: list[bytes] = []
    try:
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            pages.append(pix.tobytes("png"))
    finally:
        doc.close()
    return pages


def render_dsf_images_html(pdf_base64: str | None) -> str:
    images = pdf_base64_to_png_pages(pdf_base64)
    if not images:
        return ""

    html_parts: list[str] = []
    for image_bytes in images:
        image_b64 = base64.b64encode(image_bytes).decode("ascii")
        html_parts.append(
            f'<img src="data:image/png;base64,{image_b64}" '
            'style="width:700px;display:block;margin-bottom:10px;" />'
        )
    return "\n".join(html_parts)


def build_notification_context(
    cnpj: str,
    *,
    dsf: str,
    auditor: str,
    cargo_titulo: str = "",
    matricula: str = "",
    contato: str = "",
    orgao_origem: str = "",
    pdf_base64: str | None = None,
) -> dict[str, str]:
    normalized_cnpj = sanitize_cnpj(cnpj)
    if not validate_cnpj(normalized_cnpj):
        raise ValueError(f"CNPJ inválido: {cnpj}")

    cadastral = _load_cadastral_record(normalized_cnpj)
    malhas = _load_malhas_records(normalized_cnpj)

    razao_social = _pick_value(
        cadastral,
        "razao_social",
        "nome",
        "no_razao_social",
    )
    ie = _pick_value(cadastral, "ie", "inscricao_estadual", "co_cad_icms")

    return {
        "{{RAZAO_SOCIAL}}": razao_social,
        "{{CNPJ}}": normalized_cnpj,
        "{{IE}}": ie,
        "{{DSF}}": dsf,
        "{{AUDITOR}}": auditor,
        "{{CARGO_TITULO}}": cargo_titulo,
        "{{MATRICULA}}": matricula,
        "{{CONTATO}}": contato,
        "{{ORGAO_ORIGEM}}": orgao_origem,
        "{{TABELA}}": render_malhas_html_table(malhas),
        "{{DSF_IMAGENS}}": render_dsf_images_html(pdf_base64),
        "{{DATA_HORA_SISTEMA}}": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "{{QTD_PENDENCIAS}}": str(len(malhas)),
    }


def render_notification_txt(context: dict[str, str]) -> str:
    template_path = _template_path()
    if not template_path.exists():
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    content = template_path.read_text(encoding="utf-8")
    for placeholder, value in context.items():
        content = content.replace(placeholder, value)
    return content


def generate_notification_txt(
    cnpj: str,
    *,
    dsf: str,
    auditor: str,
    cargo_titulo: str = "",
    matricula: str = "",
    contato: str = "",
    orgao_origem: str = "",
    pdf_base64: str | None = None,
) -> tuple[str, str]:
    normalized_cnpj = sanitize_cnpj(cnpj)
    context = build_notification_context(
        normalized_cnpj,
        dsf=dsf,
        auditor=auditor,
        cargo_titulo=cargo_titulo,
        matricula=matricula,
        contato=contato,
        orgao_origem=orgao_origem,
        pdf_base64=pdf_base64,
    )
    return render_notification_txt(context), f"notificacao_det_{normalized_cnpj}.txt"


def generate_notifications_zip(
    cnpjs: list[str],
    *,
    dsf: str,
    auditor: str,
    cargo_titulo: str = "",
    matricula: str = "",
    contato: str = "",
    orgao_origem: str = "",
    pdf_base64: str | None = None,
) -> tuple[bytes, str]:
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for cnpj in cnpjs:
            content, filename = generate_notification_txt(
                cnpj,
                dsf=dsf,
                auditor=auditor,
                cargo_titulo=cargo_titulo,
                matricula=matricula,
                contato=contato,
                orgao_origem=orgao_origem,
                pdf_base64=pdf_base64,
            )
            zip_file.writestr(filename, content)

    zip_name = f"notificacoes_fisconforme_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    return zip_buffer.getvalue(), zip_name


def generate_notification_docx(
    cnpj: str,
    *,
    dsf: str,
    auditor: str,
    cargo_titulo: str = "",
    matricula: str = "",
    contato: str = "",
    orgao_origem: str = "",
    pdf_base64: str | None = None,
) -> tuple[bytes, str]:
    if Document is None or Inches is None:
        raise RuntimeError("Dependência ausente para gerar DOCX: instale `python-docx`.")

    normalized_cnpj = sanitize_cnpj(cnpj)
    if not validate_cnpj(normalized_cnpj):
        raise ValueError(f"CNPJ inválido: {cnpj}")

    cadastral = _load_cadastral_record(normalized_cnpj)
    malhas = _load_malhas_records(normalized_cnpj)

    razao_social = _pick_value(
        cadastral,
        "razao_social",
        "nome",
        "no_razao_social",
    )
    ie = _pick_value(cadastral, "ie", "inscricao_estadual", "co_cad_icms")

    document = Document()
    document.add_heading("NOTIFICAÇÃO", level=0)
    document.add_paragraph(f"Contribuinte: {razao_social}")
    document.add_paragraph(f"CNPJ: {normalized_cnpj}")
    document.add_paragraph(f"Inscrição Estadual: {ie}")
    document.add_paragraph(f"DSF: {dsf}")
    document.add_paragraph(
        "Fica o contribuinte acima identificado NOTIFICADO da existência de pendências "
        "relativas às suas obrigações tributárias acessórias registradas no sistema "
        "Fisconforme."
    )
    document.add_paragraph(
        f"Data e hora do sistema: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
    )

    document.add_heading("Pendências", level=1)
    if malhas:
        table = document.add_table(rows=1, cols=8)
        header = table.rows[0].cells
        header[0].text = "ID Pend."
        header[1].text = "ID Notif."
        header[2].text = "Malha ID"
        header[3].text = "Título"
        header[4].text = "Período"
        header[5].text = "Status Pend."
        header[6].text = "Status Notif."
        header[7].text = "Ciência"
        for malha in malhas:
            row = table.add_row().cells
            row[0].text = _pick_value(malha, "id_pendencia")
            row[1].text = _pick_value(malha, "id_notificacao")
            row[2].text = _pick_value(malha, "malhas_id")
            row[3].text = _pick_value(malha, "titulo_malha")
            row[4].text = _pick_value(malha, "periodo")
            row[5].text = _pick_value(malha, "status_pendencia")
            row[6].text = _pick_value(malha, "status_notificacao")
            row[7].text = _pick_value(malha, "data_ciencia_consolidada")
    else:
        document.add_paragraph("Sem pendências registradas.")

    images = pdf_base64_to_png_pages(pdf_base64)
    if images:
        document.add_page_break()
        document.add_heading("Imagens da DSF", level=1)
        for image_bytes in images:
            document.add_picture(BytesIO(image_bytes), width=Inches(6.5))

    document.add_page_break()
    document.add_paragraph(auditor)
    if cargo_titulo:
        document.add_paragraph(cargo_titulo)
    if matricula:
        document.add_paragraph(f"Matrícula: {matricula}")
    if contato:
        document.add_paragraph(f"Contato: {contato}")
    if orgao_origem:
        document.add_paragraph(f"Órgão de origem: {orgao_origem}")

    output = BytesIO()
    document.save(output)
    return output.getvalue(), f"notificacao_{normalized_cnpj}.docx"


def resolve_output_dir(output_dir: str) -> Path | None:
    cleaned = (output_dir or "").strip()
    if not cleaned:
        return None
    if ".." in cleaned:
        raise ValueError("Caminho não pode conter referências a diretórios superiores (..)")
    target = (settings.workspace_root / cleaned.lstrip("/\\")).resolve()
    if not str(target).startswith(str(settings.workspace_root.resolve())):
        raise ValueError("Caminho fora do diretório de trabalho permitido")
    target.mkdir(parents=True, exist_ok=True)
    return target


def save_text_output(content: str, filename: str, output_dir: str) -> str:
    target_dir = resolve_output_dir(output_dir)
    if target_dir is None:
        return ""
    target_path = target_dir / filename
    target_path.write_text(content, encoding="utf-8")
    return str(target_path)


def save_binary_output(content: bytes, filename: str, output_dir: str) -> str:
    target_dir = resolve_output_dir(output_dir)
    if target_dir is None:
        return ""
    target_path = target_dir / filename
    target_path.write_bytes(content)
    return str(target_path)
